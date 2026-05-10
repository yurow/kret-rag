"""
测试表格转 Markdown 功能
"""
import sys
from pathlib import Path

# 添加项目根目录到路径
sys.path.insert(0, str(Path(__file__).parent))

from app.services.document_service import DocumentService
from docx import Document
from openpyxl import load_workbook


def test_docx_table_conversion():
    """测试 DOCX 表格转换"""
    print("=" * 60)
    print("测试 DOCX 表格转换")
    print("=" * 60)
    
    # 创建测试文档
    test_docx = Path("test_table.docx")
    
    if not test_docx.exists():
        print(f"❌ 测试文件不存在: {test_docx}")
        print("请先上传一个包含表格的 DOCX 文件到 uploads/ 目录")
        return
    
    try:
        # 加载文档
        doc = Document(str(test_docx))
        
        print(f"\n📄 文档信息:")
        print(f"   段落数: {len(doc.paragraphs)}")
        print(f"   表格数: {len(doc.tables)}")
        
        # 创建服务实例
        service = DocumentService()
        
        # 转换每个表格
        for i, table in enumerate(doc.tables, 1):
            print(f"\n📊 表格 {i}:")
            print(f"   行数: {len(table.rows)}")
            if table.rows:
                print(f"   第一行列数: {len(table.rows[0].cells)}")
            
            # 转换为 Markdown
            markdown = service._convert_table_to_markdown(table)
            
            print(f"\n   Markdown 输出:")
            print("   " + "-" * 56)
            for line in markdown.split('\n'):
                print(f"   {line}")
            print("   " + "-" * 56)
        
        print("\n✅ DOCX 表格转换测试完成")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {str(e)}")
        import traceback
        traceback.print_exc()


def test_excel_table_conversion():
    """测试 Excel 表格转换"""
    print("\n" + "=" * 60)
    print("测试 Excel 表格转换")
    print("=" * 60)
    
    # 查找测试文件
    test_xlsx = Path("test_table.xlsx")
    
    if not test_xlsx.exists():
        print(f"❌ 测试文件不存在: {test_xlsx}")
        print("请先上传一个包含数据的 Excel 文件到 uploads/ 目录")
        return
    
    try:
        # 加载工作簿
        wb = load_workbook(str(test_xlsx), data_only=True)
        
        print(f"\n📄 工作簿信息:")
        print(f"   工作表数: {len(wb.sheetnames)}")
        print(f"   工作表名称: {', '.join(wb.sheetnames)}")
        
        # 创建服务实例
        service = DocumentService()
        
        # 转换每个工作表
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            print(f"\n📊 工作表: {sheet_name}")
            print(f"   最大行: {ws.max_row}")
            print(f"   最大列: {ws.max_column}")
            
            # 转换为 Markdown
            markdown = service._convert_excel_sheet_to_markdown(ws, sheet_name)
            
            print(f"\n   Markdown 输出:")
            print("   " + "-" * 56)
            for line in markdown.split('\n')[:20]:  # 只显示前20行
                print(f"   {line}")
            if len(markdown.split('\n')) > 20:
                print(f"   ... (共 {len(markdown.split(chr(10)))} 行)")
            print("   " + "-" * 56)
        
        print("\n✅ Excel 表格转换测试完成")
        
    except Exception as e:
        print(f"\n❌ 测试失败: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    print("表格转 Markdown 功能测试工具\n")
    
    # 运行测试
    test_docx_table_conversion()
    test_excel_table_conversion()
    
    print("\n" + "=" * 60)
    print("提示:")
    print("1. 将包含表格的文件命名为 test_table.docx 或 test_table.xlsx")
    print("2. 放置在 rag-scheduler 根目录")
    print("3. 重新运行此脚本查看转换结果")
    print("=" * 60)
